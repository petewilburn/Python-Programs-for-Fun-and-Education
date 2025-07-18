# A Python program to create a new word document

# Minimum Python version: 3.9
# Requires the python-docx library

# ----------------------------------------------------------------------------------------------
# File: P4001_Create_New_Word_Document.py
# Description: A Python program to create a new Word document using the python-docx library.
# Author: Pete W.
# License: MIT License
# Copyright (c) 2025 Pete W.
# ----------------------------------------------------------------------------------------------

from docx import Document
import os
import re

class Doc:
    def __init__(self, title: str):
        """Initialize a new document with a title."""
        if not title or not title.strip():
            raise ValueError("Title cannot be empty")
        self.title = title.strip()
        self.document = Document()
        self.document.add_heading(self.title, level=1)

    def add_paragraph(self, text: str) -> None:
        """Add a paragraph to the document."""
        if text.strip():  # Only add non-empty paragraphs
            self.document.add_paragraph(text.strip())

    def add_multiple_paragraphs(self, paragraphs: list[str]) -> None:
        """Add multiple paragraphs to the document."""
        for paragraph in paragraphs:
            self.add_paragraph(paragraph)

    def save(self, filename: str) -> None:
        """Save the document to the specified filename."""
        try:
            # Ensure the directory exists
            directory = os.path.dirname(filename)
            if directory and not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
            
            self.document.save(filename)
        except PermissionError:
            raise PermissionError(f"Permission denied: Cannot save to '{filename}'")
        except Exception as e:
            raise RuntimeError(f"Failed to save document: {e}")

class IO:
    @staticmethod
    def print_welcome_message() -> None:
        """Prints a welcome message to the user."""
        print("\nWelcome to the Word Document Creator Program!")
        print("This program will create a new Word document with your specified title and content.")
        print("You will be prompted to enter a title and content for the document.")
        print("\nLet's get started!\n")

    @staticmethod
    def prompt_for_title() -> str:
        """Prompts the user for a title and returns it."""
        count = 0
        while count < 3:
            count += 1
            title = input("Please enter the title for your document: ").strip()
            if title:  # Simplified validation - input() always returns str
                return title
            else:
                print("Invalid input. Please enter a non-empty title.")
        raise ValueError("Too many invalid attempts. Exiting the program.")
    
    @staticmethod
    def prompt_for_content() -> str:
        """Prompts the user for content and returns it."""
        count = 0
        while count < 3:
            count += 1
            content = input("Please enter the content for your document (or press Enter for empty): ").strip()
            return content  # Allow empty content
        raise ValueError("Too many invalid attempts. Exiting the program.")
    
    @staticmethod
    def prompt_for_filename() -> str:
        """Prompts the user for a filename and returns it."""
        count = 0
        while count < 3:
            count += 1
            filename = input("Please enter the filename to save your document (e.g., my_document.docx): ").strip()
            
            # Validate filename
            if not filename:
                print("Invalid input. Please enter a filename.")
                continue
            
            # Add .docx extension if missing
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            
            # Validate filename characters (Windows/cross-platform safe)
            invalid_chars = r'[<>:"/\\|?*]'
            if re.search(invalid_chars, os.path.basename(filename)):
                print("Invalid characters in filename. Avoid: < > : \" / \\ | ? *")
                continue
                
            return filename
            
        raise ValueError("Too many invalid attempts. Exiting the program.")
    
    @staticmethod
    def print_success_message(filename: str) -> None:
        """Prints a success message after saving the document."""
        abs_path = os.path.abspath(filename)
        print(f"\n✅ Success! Your document has been created and saved as:")
        print(f"   📄 {filename}")
        print(f"   📁 Full path: {abs_path}")

    @staticmethod
    def print_thank_you_message() -> None:
        """Prints a thank you message to the user."""
        print("\nThank you for using the Word Document Creator Program!")
        print("\nGoodbye!\n")

def main():
    """Main function to run the program."""
    io = IO()
    io.print_welcome_message()
    
    try:
        title = io.prompt_for_title()
        content = io.prompt_for_content()
        filename = io.prompt_for_filename()

        # Create and populate document
        doc = Doc(title)
        if content:  # Only add content if provided
            doc.add_paragraph(content)
        
        doc.save(filename)
        io.print_success_message(filename)
        
    except ValueError as e:
        print(f"❌ Input Error: {e}")
    except (PermissionError, RuntimeError) as e:
        print(f"❌ File Error: {e}")
    except ImportError:
        print("❌ Missing Dependency: Please install python-docx with: pip install python-docx")
    except Exception as e:
        print(f"❌ Unexpected Error: {e}")
    finally:
        io.print_thank_you_message()

if __name__ == "__main__":
    main()

# This program creates a new Word document with a specified title and content.
# It uses the python-docx library to handle Word document creation and manipulation.
# 
# Installation: pip install python-docx
# 
# Features:
# - Clean object-oriented design with separation of concerns
# - Robust input validation and error handling
# - Cross-platform filename validation
# - Automatic directory creation for file paths
# - Enhanced user experience with clear feedback messages
# 
# The program demonstrates:
# - Class design and encapsulation
# - Exception handling best practices  
# - Type hints for better code documentation
# - User input validation with retry logic
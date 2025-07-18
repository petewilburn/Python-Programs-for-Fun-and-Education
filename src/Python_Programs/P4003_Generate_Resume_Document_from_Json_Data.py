# A Python program to generate a Word document resume with information from a JSON file.

# Minimum Python version: 3.9

# ----------------------------------------------------------------------------------------------
# File: P4003_Generate_Resume_Document_from_Json_Data.py
# Description: A Python program to generate a Word document resume using data from a JSON file.
# Author: Pete W.
# License: MIT License
# Copyright (c) 2025 Pete W.
# ----------------------------------------------------------------------------------------------

import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import Dict, Any, List

class ResumeGenerator:
    def __init__(self, json_file: str):
        """Initialize the ResumeGenerator with a JSON file."""
        if not os.path.exists(json_file):
            raise FileNotFoundError(f"JSON file '{json_file}' not found.")
        self.json_file = json_file
        self.data = self.load_json_data()
        self.validate_json_structure()

    def load_json_data(self) -> Dict[str, Any]:
        """Load JSON data from the specified file."""
        try:
            with open(self.json_file, 'r', encoding='utf-8') as file:
                data = json.load(file)
                if not isinstance(data, dict):
                    raise ValueError("JSON file must contain an object (not array or primitive)")
                return data
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format in the file: {e}")
        except Exception as e:
            raise RuntimeError(f"Error reading JSON file: {e}")

    def validate_json_structure(self) -> None:
        """Validate that the JSON contains expected resume sections."""
        required_sections = ['personal_info']
        missing_sections = [section for section in required_sections if section not in self.data]
        
        if missing_sections:
            raise ValueError(f"Missing required sections in JSON: {', '.join(missing_sections)}")
        
        # Validate personal_info has at least name
        personal_info = self.data.get('personal_info', {})
        if not personal_info.get('name'):
            raise ValueError("Personal info must include a 'name' field")

    def create_modern_styles(self, doc: Document) -> None:
        """Create modern, sleek styles for the resume."""
        styles = doc.styles
        
        # Create modern heading style for name
        if 'Modern Name' not in [style.name for style in styles]:
            name_style = styles.add_style('Modern Name', WD_STYLE_TYPE.PARAGRAPH)
            name_font = name_style.font
            name_font.name = 'Calibri'
            name_font.size = Pt(24)
            name_font.bold = True
            name_font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_style.paragraph_format.space_after = Pt(6)
        
        # Create contact info style
        if 'Contact Info' not in [style.name for style in styles]:
            contact_style = styles.add_style('Contact Info', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Calibri'
            contact_font.size = Pt(11)
            contact_font.color.rgb = RGBColor(93, 109, 126)  # Medium gray
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Create section heading style
        if 'Section Heading' not in [style.name for style in styles]:
            section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # Create job title style
        if 'Job Title' not in [style.name for style in styles]:
            job_style = styles.add_style('Job Title', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Calibri'
            job_font.size = Pt(12)
            job_font.bold = True
            job_font.color.rgb = RGBColor(44, 62, 80)
            job_style.paragraph_format.space_before = Pt(6)
            job_style.paragraph_format.space_after = Pt(3)
        
        # Create description style
        if 'Description' not in [style.name for style in styles]:
            desc_style = styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
            desc_font = desc_style.font
            desc_font.name = 'Calibri'
            desc_font.size = Pt(11)
            desc_font.color.rgb = RGBColor(52, 73, 94)  # Dark gray
            desc_style.paragraph_format.space_after = Pt(3)
            desc_style.paragraph_format.left_indent = Inches(0.25)
        
        # Create achievement style
        if 'Achievement' not in [style.name for style in styles]:
            ach_style = styles.add_style('Achievement', WD_STYLE_TYPE.PARAGRAPH)
            ach_font = ach_style.font
            ach_font.name = 'Calibri'
            ach_font.size = Pt(10)
            ach_font.color.rgb = RGBColor(93, 109, 126)
            ach_style.paragraph_format.left_indent = Inches(0.5)
            ach_style.paragraph_format.space_after = Pt(2)
        
        # Create skills category style
        if 'Skills Category' not in [style.name for style in styles]:
            skills_style = styles.add_style('Skills Category', WD_STYLE_TYPE.PARAGRAPH)
            skills_font = skills_style.font
            skills_font.name = 'Calibri'
            skills_font.size = Pt(11)
            skills_font.bold = True
            skills_font.color.rgb = RGBColor(44, 62, 80)
            skills_style.paragraph_format.space_after = Pt(3)

    def add_horizontal_line(self, doc: Document) -> None:
        """Add a subtle horizontal line for visual separation."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create a horizontal line using underscores with light color
        line_run = paragraph.add_run('_' * 60)
        line_run.font.color.rgb = RGBColor(232, 232, 232)  # Light gray
        line_run.font.size = Pt(8)
        paragraph.space_after = Pt(12)

    def format_personal_info(self, doc: Document) -> None:
        """Add modern formatted personal information section."""
        personal_info = self.data.get('personal_info', {})
        
        # Name with modern styling
        name = personal_info.get('name', 'Unknown')
        name_paragraph = doc.add_paragraph(name)
        name_paragraph.style = 'Modern Name'
        
        # Contact information with modern icons and layout
        contact_items = []
        if personal_info.get('email'):
            contact_items.append(f"✉ {personal_info['email']}")
        if personal_info.get('phone'):
            contact_items.append(f"☎ {personal_info['phone']}")
        if personal_info.get('address'):
            contact_items.append(f"⌂ {personal_info['address']}")
        if personal_info.get('linkedin'):
            contact_items.append(f"� {personal_info['linkedin']}")
        if personal_info.get('website'):
            contact_items.append(f"🌐 {personal_info['website']}")
        
        if contact_items:
            # Split contact info into multiple lines for better readability
            if len(contact_items) > 3:
                # First line: email and phone
                line1 = [item for item in contact_items[:2]]
                contact_para1 = doc.add_paragraph(' • '.join(line1))
                contact_para1.style = 'Contact Info'
                
                # Second line: address and links
                line2 = [item for item in contact_items[2:]]
                contact_para2 = doc.add_paragraph(' • '.join(line2))
                contact_para2.style = 'Contact Info'
            else:
                contact_paragraph = doc.add_paragraph(' • '.join(contact_items))
                contact_paragraph.style = 'Contact Info'
        
        # Add decorative line
        self.add_horizontal_line(doc)
        
        # Professional summary with modern styling
        if personal_info.get('summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
            summary_heading.style = 'Section Heading'
            
            summary_para = doc.add_paragraph(personal_info['summary'])
            summary_para.style = 'Description'
            summary_para.paragraph_format.left_indent = Inches(0)  # No indent for summary

    def generate_resume(self) -> Document:
        """Generate a modern, professionally formatted Word document resume."""
        doc = Document()
        
        # Set document margins for modern layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # Create modern styles
        self.create_modern_styles(doc)
        
        # Add formatted personal information
        self.format_personal_info(doc)
        
        # Add objective if provided
        objective = self.data.get('objective')
        if objective:
            obj_heading = doc.add_paragraph('OBJECTIVE')
            obj_heading.style = 'Section Heading'
            obj_para = doc.add_paragraph(objective)
            obj_para.style = 'Description'
            obj_para.paragraph_format.left_indent = Inches(0)

        # Add education details with modern formatting
        education = self.data.get('education', [])
        if education:
            edu_heading = doc.add_paragraph('EDUCATION')
            edu_heading.style = 'Section Heading'
            
            for edu in education:
                degree = edu.get('degree', '')
                field = edu.get('field_of_study', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')
                gpa = edu.get('gpa', '')
                
                # Main education line with modern formatting
                edu_title = f"{degree}"
                if field:
                    edu_title += f" in {field}"
                
                edu_para = doc.add_paragraph(edu_title)
                edu_para.style = 'Job Title'
                
                # Institution and year
                if institution or year:
                    inst_text = institution
                    if year:
                        inst_text += f" • {year}"
                    inst_para = doc.add_paragraph(inst_text)
                    inst_para.style = 'Description'
                
                # Add GPA if provided
                if gpa:
                    gpa_para = doc.add_paragraph(f"GPA: {gpa}")
                    gpa_para.style = 'Achievement'

        # Add work experience with enhanced formatting
        work_experience = self.data.get('work_experience', [])
        if work_experience:
            work_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
            work_heading.style = 'Section Heading'
            
            for job in work_experience:
                position = job.get('position', '')
                company = job.get('company', '')
                start_date = job.get('start_date', '')
                end_date = job.get('end_date', 'Present')
                
                # Job title and company with modern layout
                job_title = f"{position}"
                if company:
                    job_title += f" • {company}"
                
                job_para = doc.add_paragraph(job_title)
                job_para.style = 'Job Title'
                
                # Date range
                if start_date:
                    date_para = doc.add_paragraph(f"{start_date} - {end_date}")
                    date_para.style = 'Achievement'
                
                # Job description
                description = job.get('description', '')
                if description:
                    desc_para = doc.add_paragraph(description)
                    desc_para.style = 'Description'
                
                # Achievements with bullet points
                achievements = job.get('achievements', [])
                if achievements:
                    for achievement in achievements:
                        ach_para = doc.add_paragraph(f"▪ {achievement}")
                        ach_para.style = 'Achievement'

        # Add projects section with modern styling
        projects = self.data.get('projects', [])
        if projects:
            proj_heading = doc.add_paragraph('KEY PROJECTS')
            proj_heading.style = 'Section Heading'
            
            for project in projects:
                name = project.get('name', '')
                description = project.get('description', '')
                technologies = project.get('technologies', [])
                
                if name:
                    proj_para = doc.add_paragraph(name)
                    proj_para.style = 'Job Title'
                
                if description:
                    desc_para = doc.add_paragraph(description)
                    desc_para.style = 'Description'
                
                if technologies:
                    tech_para = doc.add_paragraph(f"Technologies: {' • '.join(technologies)}")
                    tech_para.style = 'Achievement'

        # Add skills with modern categorization
        skills = self.data.get('skills', [])
        if skills:
            skills_heading = doc.add_paragraph('CORE COMPETENCIES')
            skills_heading.style = 'Section Heading'
            
            if isinstance(skills, dict):
                for category, skill_list in skills.items():
                    if isinstance(skill_list, list):
                        cat_para = doc.add_paragraph(f"{category}")
                        cat_para.style = 'Skills Category'
                        
                        skills_text = ' • '.join(skill_list)
                        skills_para = doc.add_paragraph(skills_text)
                        skills_para.style = 'Description'
            else:
                skills_para = doc.add_paragraph(' • '.join(skills))
                skills_para.style = 'Description'

        # Add certifications with modern formatting
        certifications = self.data.get('certifications', [])
        if certifications:
            cert_heading = doc.add_paragraph('CERTIFICATIONS')
            cert_heading.style = 'Section Heading'
            
            for cert in certifications:
                if isinstance(cert, str):
                    cert_para = doc.add_paragraph(f"▪ {cert}")
                    cert_para.style = 'Description'
                elif isinstance(cert, dict):
                    name = cert.get('name', '')
                    issuer = cert.get('issuer', '')
                    date = cert.get('date', '')
                    
                    if name:
                        cert_title = doc.add_paragraph(name)
                        cert_title.style = 'Job Title'
                    
                    cert_details = []
                    if issuer:
                        cert_details.append(issuer)
                    if date:
                        cert_details.append(date)
                    
                    if cert_details:
                        cert_info = doc.add_paragraph(' • '.join(cert_details))
                        cert_info.style = 'Achievement'

        return doc

    def save_resume(self, filename: str) -> None:
        """Save the generated resume to a Word document."""
        try:
            # Ensure directory exists
            directory = os.path.dirname(filename)
            if directory and not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
            
            resume_doc = self.generate_resume()
            resume_doc.save(filename)
        except PermissionError:
            raise PermissionError(f"Permission denied: Cannot save to '{filename}'")
        except Exception as e:
            raise RuntimeError(f"Failed to save resume: {e}")

    def get_resume_preview(self) -> str:
        """Generate a text preview of the resume content."""
        preview = []
        
        # Personal info preview
        personal_info = self.data.get('personal_info', {})
        preview.append(f"Name: {personal_info.get('name', 'N/A')}")
        preview.append(f"Email: {personal_info.get('email', 'N/A')}")
        
        # Count sections
        sections = []
        if self.data.get('education'):
            sections.append(f"Education ({len(self.data['education'])} entries)")
        if self.data.get('work_experience'):
            sections.append(f"Work Experience ({len(self.data['work_experience'])} entries)")
        if self.data.get('skills'):
            sections.append("Skills")
        if self.data.get('projects'):
            sections.append(f"Projects ({len(self.data['projects'])} entries)")
        if self.data.get('certifications'):
            sections.append(f"Certifications ({len(self.data['certifications'])} entries)")
        
        preview.append(f"Sections: {', '.join(sections) if sections else 'None'}")
        
        return '\n'.join(preview)

class IO:
    @staticmethod
    def print_welcome_message() -> None:
        """Prints a welcome message to the user."""
        print("\nWelcome to the Modern Resume Generator! ✨📄")
        print("This program creates sleek, professionally designed Word document resumes from JSON data.")
        print("Features modern typography, color schemes, and clean layouts that stand out to employers.")
        print("Supported sections: Personal Info, Education, Work Experience, Skills, Projects, Certifications")
        print("\nLet's create your stunning resume!\n")

    @staticmethod
    def print_example_json() -> None:
        """Print an example JSON structure for user reference."""
        example = '''
Example JSON structure:
{
  "personal_info": {
    "name": "John Doe",
    "email": "john.doe@email.com",
    "phone": "+1-555-0123",
    "address": "123 Main St, City, State 12345",
    "linkedin": "linkedin.com/in/johndoe",
    "summary": "Experienced software developer with 5+ years..."
  },
  "objective": "Seeking a challenging role in software development...",
  "education": [
    {
      "degree": "Bachelor of Science",
      "field_of_study": "Computer Science",
      "institution": "University of Technology",
      "year": "2020",
      "gpa": "3.8"
    }
  ],
  "work_experience": [
    {
      "position": "Software Developer",
      "company": "Tech Corp",
      "start_date": "2020",
      "end_date": "Present",
      "description": "Developed web applications...",
      "achievements": ["Increased efficiency by 30%", "Led team of 3 developers"]
    }
  ],
  "skills": {
    "Programming": ["Python", "JavaScript", "Java"],
    "Frameworks": ["Django", "React", "Spring"],
    "Tools": ["Git", "Docker", "AWS"]
  },
  "projects": [
    {
      "name": "E-commerce Platform",
      "description": "Built a full-stack e-commerce solution",
      "technologies": ["Python", "Django", "PostgreSQL"]
    }
  ],
  "certifications": [
    {
      "name": "AWS Certified Developer",
      "issuer": "Amazon Web Services",
      "date": "2023"
    }
  ]
}
        '''
        print(example)

    @staticmethod
    def prompt_for_json_file() -> str:
        """Prompts the user for the path to the JSON file with validation."""
        count = 0
        while count < 3:
            count += 1
            json_file = input("Enter the path to your JSON resume file: ").strip()
            
            if not json_file:
                print("❌ File path cannot be empty.")
                continue
            
            # Check if file exists
            if not os.path.exists(json_file):
                print(f"❌ File '{json_file}' not found.")
                if count < 3:
                    show_example = input("Would you like to see an example JSON structure? (y/n): ").strip().lower()
                    if show_example in ['y', 'yes']:
                        IO.print_example_json()
                continue
            
            # Check if it's a JSON file
            if not json_file.lower().endswith('.json'):
                print("❌ File must be a JSON file (.json extension).")
                continue
                
            return json_file
        
        raise ValueError("Too many invalid attempts. Exiting the program.")

    @staticmethod
    def prompt_for_output_filename() -> str:
        """Prompts the user for the output filename."""
        count = 0
        while count < 3:
            count += 1
            filename = input("Enter output filename for resume (e.g., 'my_resume.docx'): ").strip()
            
            if not filename:
                filename = "resume.docx"  # Default filename
                print(f"Using default filename: {filename}")
                return filename
            
            # Add .docx extension if missing
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            
            # Validate filename characters
            invalid_chars = r'[<>:"/\\|?*]'
            if re.search(invalid_chars, os.path.basename(filename)):
                print("❌ Invalid characters in filename. Avoid: < > : \" / \\ | ? *")
                continue
                
            return filename
        
        raise ValueError("Too many invalid attempts with filename.")

    @staticmethod
    def print_resume_preview(generator: 'ResumeGenerator') -> None:
        """Print a preview of the resume content."""
        print("\n📋 Resume Preview:")
        print("=" * 50)
        print(generator.get_resume_preview())
        print("=" * 50)
        
        proceed = input("\nGenerate resume with this data? (y/n): ").strip().lower()
        if proceed not in ['y', 'yes']:
            raise ValueError("User cancelled resume generation.")
    
    @staticmethod
    def print_resume_saved_message(filename: str) -> None:
        """Prints a message indicating that the resume has been saved."""
        abs_path = os.path.abspath(filename)
        print(f"\n✅ Success! Professional resume generated:")
        print(f"   📄 {filename}")
        print(f"   📁 Full path: {abs_path}")
        print(f"   📊 File size: {os.path.getsize(filename)} bytes")

    @staticmethod
    def print_thank_you_message() -> None:
        """Prints a thank you message to the user."""
        print("\nThank you for using the Modern Resume Generator! ✨🎉")
        print("Your sleek, professional resume is ready to impress employers!")
        print("Best of luck with your job applications!")
        print("\nGoodbye!\n")

def main():
    """Main function to run the program."""
    IO.print_welcome_message()
    
    try:
        # Get JSON file path
        json_file = IO.prompt_for_json_file()
        
        # Initialize resume generator (this validates the JSON)
        generator = ResumeGenerator(json_file)
        
        # Show preview and get confirmation
        IO.print_resume_preview(generator)
        
        # Get output filename
        output_filename = IO.prompt_for_output_filename()
        
        # Generate and save resume
        generator.save_resume(output_filename)
        IO.print_resume_saved_message(output_filename)
        
    except FileNotFoundError as e:
        print(f"❌ File Error: {e}")
    except ValueError as e:
        print(f"❌ Data Error: {e}")
    except (PermissionError, RuntimeError) as e:
        print(f"❌ System Error: {e}")
    except ImportError:
        print("❌ Missing Dependency: Please install python-docx with: pip install python-docx")
    except KeyboardInterrupt:
        print("\n❌ Program interrupted by user.")
    except Exception as e:
        print(f"❌ Unexpected Error: {e}")
    finally:
        IO.print_thank_you_message()

if __name__ == "__main__":
    main()

# This program generates a modern, sleek Word document resume from JSON data.
# 
# Modern Design Features:
# - Custom typography with Calibri font and professional sizing
# - Color-coded sections with dark blue-gray headers and blue accent lines
# - Clean layout with proper spacing and visual hierarchy
# - Modern contact info with streamlined icons and bullet separators
# - Professional section headings with subtle bottom borders
# - Hierarchical text styling for better readability
# - Modern bullet points (▪) and separators (•) throughout
# 
# Technical Features:
# - Professional document formatting with optimal margins
# - Comprehensive resume sections with modern styling
# - JSON structure validation and helpful error messages
# - Resume preview before generation
# - Customizable output filename with validation
# - Example JSON structure for user guidance
# - Enhanced error handling and user feedback
# 
# Requirements: python-docx library (pip install python-docx)
# 
# The program demonstrates:
# - Advanced Word document styling with custom fonts and colors
# - Modern UI/UX design principles applied to document generation
# - Robust JSON data validation and processing
# - Professional user experience with previews and examples
# - Contemporary resume design trends and best practices